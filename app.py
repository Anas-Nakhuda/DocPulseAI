import os
import fitz
import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from docx import Document as DocxDocument

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.documents import Document

load_dotenv()

st.set_page_config(page_title="DocPulse AI — Document Intelligence", page_icon="🩺", layout="wide")

INDEX_DIR = "faiss_index"
DEFAULT_CHAT_MODEL = os.getenv("GEMINI_CHAT_MODEL", "gemini-3.1-flash-lite")

# Initialize lightweight local embeddings
@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

embeddings = get_embeddings()


def check_model_availability():
    """Ping Google's model list once per session to catch deprecated model names early."""
    if "model_check_done" in st.session_state:
        return
    st.session_state.model_check_done = True
    try:
        from google import genai
        client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
        available = {m.name.split("/")[-1] for m in client.models.list()}
        if DEFAULT_CHAT_MODEL not in available:
            st.session_state.model_warning = (
                f"⚠️ Configured model '{DEFAULT_CHAT_MODEL}' not found on your API key. "
                f"Google may have deprecated it. Set GEMINI_CHAT_MODEL in your .env file."
            )
    except Exception:
        pass


def inject_theme():
    css_rules = """
:root {
    --dp-bg: #14161C;
    --dp-surface: #1E212B;
    --dp-surface-2: #262A36;
    --dp-border: #33384A;
    --dp-text: #EDEDF0;
    --dp-text-muted: #8A8D98;
    --dp-amber: #E8A33D;
    --dp-teal: #4FD1C5;
}
.stApp { background-color: var(--dp-bg); }
html, body, [class*="css"] { font-family: 'Inter', sans-serif; color: var(--dp-text); }
h1, h2, h3, .dp-brand { font-family: 'Space Grotesk', sans-serif !important; }
section[data-testid="stSidebar"] { background-color: var(--dp-surface); border-right: 1px solid var(--dp-border); }
.stButton>button[kind="primary"] { background-color: var(--dp-amber); color: #14161C; border: none; font-weight: 600; }
.stButton>button[kind="primary"]:hover { background-color: #f2b45a; }
[data-testid="stChatMessage"] { background-color: var(--dp-surface); border: 1px solid var(--dp-border); border-radius: 12px; padding: 0.25rem 0.5rem; }
.dp-header { padding: 1.25rem 0 0.5rem 0; border-bottom: 1px solid var(--dp-border); margin-bottom: 1.25rem; }
.dp-brand { font-size: 1.9rem; font-weight: 700; color: var(--dp-text); margin: 0; letter-spacing: -0.02em; }
.dp-brand span { color: var(--dp-amber); }
.dp-tagline { color: var(--dp-text-muted); font-size: 0.95rem; margin-top: 0.15rem; }
.dp-pulse-wrap { margin-top: 0.6rem; height: 28px; }
.dp-pulse-wrap svg { width: 100%; height: 28px; }
.dp-pulse-line { fill: none; stroke: var(--dp-amber); stroke-width: 2; stroke-linecap: round; stroke-linejoin: round; stroke-dasharray: 340; stroke-dashoffset: 340; animation: dp-draw 2.6s ease-in-out infinite; opacity: 0.85; }
@keyframes dp-draw { 0% { stroke-dashoffset: 340; } 50% { stroke-dashoffset: 0; } 100% { stroke-dashoffset: -340; } }
.dp-chip { display: inline-flex; align-items: center; gap: 0.4rem; background-color: var(--dp-surface-2); border: 1px solid var(--dp-border); border-radius: 999px; padding: 0.3rem 0.75rem; margin: 0.2rem 0.3rem 0.2rem 0; font-family: 'JetBrains Mono', monospace; font-size: 0.78rem; color: var(--dp-teal); }
.dp-chip b { color: var(--dp-text); font-weight: 600; }
.dp-chunk-preview { font-family: 'JetBrains Mono', monospace; font-size: 0.8rem; color: var(--dp-text-muted); background-color: var(--dp-surface-2); border-left: 2px solid var(--dp-teal); padding: 0.5rem 0.75rem; margin: 0.3rem 0 0.9rem 0; border-radius: 0 6px 6px 0; }
.dp-status { font-family: 'JetBrains Mono', monospace; font-size: 0.8rem; padding: 0.3rem 0.6rem; border-radius: 6px; display: inline-block; margin-top: 0.4rem; }
.dp-status-ready { background-color: rgba(79,209,197,0.12); color: var(--dp-teal); border: 1px solid rgba(79,209,197,0.3); }
.dp-status-empty { background-color: rgba(138,141,152,0.12); color: var(--dp-text-muted); border: 1px solid var(--dp-border); }
""".strip()

    injector = f"""
    <script>
    (function() {{
        const doc = window.parent.document;
        if (!doc.getElementById('dp-theme-style')) {{
            const style = doc.createElement('style');
            style.id = 'dp-theme-style';
            style.textContent = {css_rules!r};
            doc.head.appendChild(style);
        }}
        if (!doc.getElementById('dp-fonts')) {{
            const link1 = doc.createElement('link');
            link1.rel = 'preconnect';
            link1.href = 'https://fonts.googleapis.com';
            const link2 = doc.createElement('link');
            link2.id = 'dp-fonts';
            link2.rel = 'stylesheet';
            link2.href = 'https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;700&family=Inter:wght@400;500;600&family=JetBrains+Mono:wght@400;600&display=swap';
            doc.head.appendChild(link1);
            doc.head.appendChild(link2);
        }}
    }})();
    </script>
    """
    components.html(injector, height=0, width=0)


def render_header(active=False):
    animation_state = "running" if active else "paused"
    html = f"""
        <div class="dp-header">
            <p class="dp-brand">Doc<span>Pulse</span> AI</p>
            <p class="dp-tagline">Ask questions over your documents — with cited, traceable answers.</p>
            <div class="dp-pulse-wrap">
                <svg viewBox="0 0 340 28" preserveAspectRatio="none" style="animation-play-state:{animation_state};">
                    <path class="dp-pulse-line" style="animation-play-state:{animation_state};"
                        d="M0,14 L60,14 L75,4 L90,24 L105,14 L130,14 L145,20 L160,8 L175,14 L340,14" />
                </svg>
            </div>
        </div>
        """
    st.markdown("\n".join(line.strip() for line in html.strip().split("\n")), unsafe_allow_html=True)


def render_citation_chips(results):
    st.markdown('<p style="color:var(--dp-text-muted); font-size:0.82rem; margin:0.4rem 0 0.3rem 0;">Answered using:</p>', unsafe_allow_html=True)
    seen = {}
    for doc, score in results:
        key = (doc.metadata.get("source"), doc.metadata.get("page"))
        if key not in seen or score < seen[key]:
            seen[key] = score

    chips_html = "".join(f'<span class="dp-chip"><b>{source}</b> · page {page}</span>' for (source, page) in seen.keys())
    st.markdown(chips_html, unsafe_allow_html=True)


def extract_documents(uploaded_files, progress_callback=None):
    docs = []
    for file in uploaded_files:
        name = file.name
        if name.lower().endswith(".pdf"):
            pdf_bytes = file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            total_pages = len(doc)
            for page_num in range(1, total_pages + 1):
                page = doc[page_num - 1]
                text = page.get_text() or ""
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"source": name, "page": page_num}))
                if progress_callback and (page_num % 5 == 0 or page_num == total_pages):
                    progress_callback(name, page_num, total_pages)
        elif name.lower().endswith(".docx"):
            if progress_callback:
                progress_callback(name, 0, 1)
            docx_file = DocxDocument(file)
            text = "\n".join(p.text for p in docx_file.paragraphs if p.text.strip())
            if text.strip():
                docs.append(Document(page_content=text, metadata={"source": name, "page": "N/A"}))
        elif name.lower().endswith(".txt"):
            if progress_callback:
                progress_callback(name, 0, 1)
            text = file.read().decode("utf-8", errors="ignore")
            if text.strip():
                docs.append(Document(page_content=text, metadata={"source": name, "page": "N/A"}))
        else:
            st.warning(f"Skipped unsupported file type: {name}")
    return docs


def chunk_documents(docs, chunk_size=1000, chunk_overlap=200):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(docs)


def build_vector_store(chunks, progress_callback=None):
    if progress_callback:
        progress_callback(0, len(chunks), "Generating local embeddings for vector store...")
    
    vector_store = FAISS.from_documents(chunks, embeddings)
    vector_store.save_local(INDEX_DIR)
    st.session_state.vector_store = vector_store
    
    if progress_callback:
        progress_callback(len(chunks), len(chunks), "Vector store built and saved successfully!")
        
    return vector_store


CONTEXTUALIZE_PROMPT = PromptTemplate.from_template(
    """Given the recent chat history and a new question, rewrite the new
question as a standalone question that contains all necessary context.
If it is already standalone, return it unchanged. Do not answer it.

Chat history:
{chat_history}

New question: {question}

Standalone question:"""
)

ANSWER_PROMPT = PromptTemplate.from_template(
    """Answer the question as precisely and completely as possible using ONLY
the provided context. If the answer is not present in the context, say
"The answer is not available in the uploaded document(s)." Do not guess.

Context:
{context}

Question: {question}

Answer:"""
)


def format_chat_history(messages, max_turns=3):
    recent = messages[-(max_turns * 2):]
    return "\n".join(f"{m['role']}: {m['content']}" for m in recent) or "None"


def answer_question(user_question, model_name, temperature, top_k):
    vector_store = None
    if "vector_store" in st.session_state and st.session_state.vector_store is not None:
        vector_store = st.session_state.vector_store
    elif os.path.exists(INDEX_DIR):
        vector_store = FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)
    
    if vector_store is None:
        return "No indexed content found. Please upload and process a document first.", []

    model = ChatGoogleGenerativeAI(model=model_name, temperature=temperature)

    history_text = format_chat_history(st.session_state.get("messages", []))
    if history_text != "None":
        standalone_chain = CONTEXTUALIZE_PROMPT | model | StrOutputParser()
        standalone_question = standalone_chain.invoke(
            {"chat_history": history_text, "question": user_question}
        ).strip()
    else:
        standalone_question = user_question

    results = vector_store.similarity_search_with_score(standalone_question, k=top_k)

    if not results:
        return "No relevant information found in the processed documents.", []

    context_text = "\n\n".join(doc.page_content for doc, _ in results)
    answer_chain = ANSWER_PROMPT | model | StrOutputParser()
    answer = answer_chain.invoke({"context": context_text, "question": standalone_question})

    return answer, results


# ---------------------------------------------------------------------------
# Streamlit Interface
# ---------------------------------------------------------------------------
inject_theme()

if "messages" not in st.session_state:
    st.session_state.messages = []
if "index_ready" not in st.session_state:
    st.session_state.index_ready = os.path.exists(INDEX_DIR)

check_model_availability()
render_header(active=True)

if st.session_state.get("model_warning"):
    st.warning(st.session_state.model_warning)

with st.sidebar:
    st.markdown('<p class="dp-brand" style="font-size:1.1rem;">📁 Documents</p>', unsafe_allow_html=True)

    if st.session_state.index_ready:
        st.markdown('<span class="dp-status dp-status-ready">● Index ready — ask away</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="dp-status dp-status-empty">○ No document indexed yet</span>', unsafe_allow_html=True)

    st.write("")
    uploaded_files = st.file_uploader("Upload PDF, DOCX, or TXT files", accept_multiple_files=True, type=["pdf", "docx", "txt"])

    with st.expander("⚙️ Advanced settings"):
        model_name = st.selectbox(
            "Model",
            options=["gemini-3.1-flash-lite", "gemini-3-flash-preview", "gemini-3.1-pro-preview"],
            index=0
        )
        temperature = st.slider("Temperature", 0.0, 1.0, 0.3, 0.1)
        chunk_size = st.slider("Chunk size", 500, 2000, 1000, 100)
        top_k = st.slider("Chunks retrieved (k)", 2, 8, 4, 1)

    if st.button("Process Documents", type="primary"):
        if not uploaded_files:
            st.error("Please upload at least one file first.")
        elif not os.getenv("GOOGLE_API_KEY"):
            st.error("GOOGLE_API_KEY not found. Add it to your .env file.")
        else:
            try:
                extraction_status = st.empty()
                extraction_status.markdown("📄 **Reading your file(s)...**")

                def on_extract_progress(filename, done, total):
                    if total > 1:
                        extraction_status.markdown(f"📄 **Reading {filename}** — page {done}/{total}...")

                docs = extract_documents(uploaded_files, progress_callback=on_extract_progress)
                extraction_status.empty()

                if not docs:
                    st.error("Couldn't extract any text. Are the PDFs scanned images without OCR?")
                else:
                    chunks = chunk_documents(docs, chunk_size=chunk_size)
                    progress_bar = st.progress(0.0)
                    status_text = st.empty()

                    def on_progress(done, total, message):
                        progress_bar.progress(min(done / total, 1.0))
                        status_text.markdown("⏳ Processing vector embeddings...")

                    on_progress(0, len(chunks), None)
                    build_vector_store(chunks, progress_callback=on_progress)

                    progress_bar.empty()
                    status_text.empty()
                    st.session_state.index_ready = True
                    st.session_state.messages = []
                    st.success(f"Indexed {len(docs)} sections into {len(chunks)} chunks!")
            except Exception as e:
                st.error("Something went wrong while processing this document.")
                with st.expander("Technical details"):
                    st.code(str(e)[:1000])

    if st.session_state.index_ready and st.button("🗑️ Clear index"):
        import shutil
        shutil.rmtree(INDEX_DIR, ignore_errors=True)
        st.session_state.vector_store = None
        st.session_state.index_ready = False
        st.session_state.messages = []
        st.rerun()

if not st.session_state.messages:
    empty_title = "Ready when you are" if st.session_state.index_ready else "Nothing indexed yet"
    empty_body = "Ask anything about your document." if st.session_state.index_ready else "Upload a document to start."
    st.markdown(f"""
        <div style="text-align:center; padding: 3rem 1rem; color: var(--dp-text-muted);">
            <p style="font-family:'Space Grotesk',sans-serif; font-size:1.1rem; color: var(--dp-text);">{empty_title}</p>
            <p style="font-size:0.9rem;">{empty_body}</p>
        </div>
    """, unsafe_allow_html=True)

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

user_question = st.chat_input("Ask a question about your document(s)...")

if user_question:
    if not st.session_state.index_ready:
        st.error("No document indexed yet. Upload and process a file from the sidebar first.")
    else:
        with st.chat_message("user"):
            st.write(user_question)
        st.session_state.messages.append({"role": "user", "content": user_question})

        with st.chat_message("assistant"):
            with st.spinner("Searching document and generating answer..."):
                try:
                    answer, results = answer_question(user_question, model_name, temperature, top_k)
                    st.write(answer)

                    if results:
                        render_citation_chips(results)
                        with st.expander("🔍 View retrieved chunks (transparency panel)"):
                            for i, (doc, score) in enumerate(results, start=1):
                                st.markdown(f"**[{i}] {doc.metadata.get('source')} · p.{doc.metadata.get('page')} · distance: {score:.3f}**")
                                st.markdown(f'<div class="dp-chunk-preview">{doc.page_content[:300]}...</div>', unsafe_allow_html=True)

                    st.session_state.messages.append({"role": "assistant", "content": answer})
                except Exception as e:
                    st.error("Something went wrong while generating an answer.")
                    with st.expander("Technical details"):
                        st.code(str(e)[:1000])

if st.session_state.messages:
    transcript = "\n\n".join(f"{m['role'].upper()}: {m['content']}" for m in st.session_state.messages)
    st.sidebar.download_button("⬇️ Download chat transcript", transcript, file_name="docpulse_chat.txt")
